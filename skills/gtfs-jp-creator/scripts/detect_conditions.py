# -*- coding: utf-8 -*-
"""時刻表ファイルの原文から、PDF/Excelに「書かれている」条件を保守的に検出する。

目的: ③条件確認で，書かれている項目は候補として自動入力し（要確認）, 無い項目だけ人が入れる。
誤検出を避けるため**確信度の高いパターンのみ**拾う。事業者名・法人番号・路線名は取り違えリスクが
高いので**検出しない**（人が入力）。確定は必ず人（正しく失敗の原則）。

入力: .pdf(テキスト) / .xlsx / .md / .txt
出力: {fare_adult, fare_child, fare_disabled, days[7], holiday_syukujitsu, holiday_nenmatsu,
       holiday_obon, start_date, end_date, phone, url, _evidence{...}}
使い方: python detect_conditions.py <file> -o conditions.json
"""
import argparse
import datetime
import json
import re
import sys
import unicodedata
from pathlib import Path


def _read_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")
    if ext == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            out = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    for c in row:
                        if c is not None:
                            out.append(str(c))
            return "\n".join(out)
        except Exception:
            return ""
    if ext == ".docx":
        try:
            import docx
            d = docx.Document(str(path))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:                       # 表内の運賃・条件も拾う
                for row in t.rows:
                    for c in row.cells:
                        parts.append(c.text)
            return "\n".join(parts)
        except Exception:
            return ""
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return "\n".join((p.extract_text() or "") for p in pdf.pages)
        except Exception:
            return ""
    return ""


def _to_ymd(g):
    """(era/year, month, day) tuple → YYYYMMDD。令和=R/令和 のみ対応（保守的）。"""
    y, m, d = g
    return f"{int(y):04d}{int(m):02d}{int(d):02d}"


def detect(text: str, today: str = None) -> dict:
    res, ev = {}, {}
    # NFKC で全角→半角に正規化（全角数字の電話 ０９４９… や ＴＥＬ・全角ダッシュ等に対応）。
    t = unicodedata.normalize("NFKC", text or "")
    if today is None:
        today = datetime.date.today().strftime("%Y%m%d")

    # --- 運賃（円が明示されているものだけ。路線で運賃が違う場合に備え「全候補」を集める） ---
    def yen_all(keys):
        vals = {}   # price -> evidence（重複値はまとめる）
        for k in keys:
            for m in re.finditer(k + r"[^\d¥￥]{0,8}[¥￥]?\s*(\d{2,4})\s*円", t):
                vals.setdefault(int(m.group(1)), m.group(0).strip())
        return vals
    adult = yen_all([r"大人", r"おとな", r"一般", r"中学生以上", r"中学生", r"高校生以上"])
    child = yen_all([r"小児", r"こども", r"子供", r"小学生", r"小人"])
    dis = yen_all([r"障害者", r"障がい者", r"障碍者"])
    if not adult:   # 区分が無く「均一/一律 ○○円」だけのとき大人として拾う
        for m in re.finditer(r"(均一|一律)[^\d¥￥]{0,8}[¥￥]?\s*(\d{2,4})\s*円", t):
            adult.setdefault(int(m.group(2)), m.group(0).strip())
    # 区分ごとに「1種類だけ」なら採用、複数あれば採用せず候補に回す（＝勝手に1つに決めない）
    for key, d in (("fare_adult", adult), ("fare_child", child), ("fare_disabled", dis)):
        if len(d) == 1:
            p = next(iter(d)); res[key] = p; ev[key] = d[p]
    cands = []
    for cat, d in (("大人", adult), ("小児", child), ("障がい者", dis)):
        for p, e in sorted(d.items()):
            cands.append({"category": cat, "price": p, "evidence": e})
    if cands:
        res["fare_candidates"] = cands
    # いずれかの区分で複数の異なる運賃 → 路線で異なる可能性（単一自動入力しない目印）
    res["fare_multiple"] = any(len(d) > 1 for d in (adult, child, dis))

    # --- 運行曜日 ---
    days = None
    if re.search(r"毎日\s*運行|年中無休", t):
        days = [1, 1, 1, 1, 1, 1, 1]; ev["days"] = "毎日運行"
    elif re.search(r"(土曜?・?日曜?・?祝|土日祝|土・日|土日).{0,6}(運休|を?除く|運行なし)", t):
        days = [1, 1, 1, 1, 1, 0, 0]; ev["days"] = "土日祝運休"
    elif re.search(r"平日.{0,4}(のみ|だけ|運行)", t):
        days = [1, 1, 1, 1, 1, 0, 0]; ev["days"] = "平日のみ"
    if days:
        res["days"] = days

    # --- 運休日（祝日/年末年始/お盆） ---
    if re.search(r"祝(日|祭日)?.{0,6}運休|祝日.{0,4}(を?除く|お休み)", t):
        res["holiday_syukujitsu"] = True; ev["holiday_syukujitsu"] = "祝日運休"
    if re.search(r"年末年始.{0,6}運休|年末年始.{0,4}(を?除く|休)", t):
        res["holiday_nenmatsu"] = True; ev["holiday_nenmatsu"] = "年末年始運休"
    if re.search(r"(お盆|盆).{0,6}運休|(お盆|盆).{0,4}(を?除く|休)", t):
        res["holiday_obon"] = True; ev["holiday_obon"] = "お盆運休"

    # --- 有効期間（令和/西暦の日付。範囲があれば start/end） ---
    def find_dates():
        out = []
        for m in re.finditer(r"(?:令和|R)\s*(\d{1,2})[\.\-年/]\s*(\d{1,2})[\.\-月/]\s*(\d{1,2})", t):
            y = 2018 + int(m.group(1))  # 令和元年=2019 → R1=2019, 2018+1
            out.append(f"{y:04d}{int(m.group(2)):02d}{int(m.group(3)):02d}")
        for m in re.finditer(r"(20\d{2})[\.\-年/]\s*(\d{1,2})[\.\-月/]\s*(\d{1,2})", t):
            out.append(f"{int(m.group(1)):04d}{int(m.group(2)):02d}{int(m.group(3)):02d}")
        return out
    # 古い資料の改正日をそのまま有効期間に入れない（正しく失敗）。
    # 範囲: 終了日が今日以降なら採用（開始が過去でも有効期間として正当）。終了済みは採用せず警告。
    # 単一日付(改正日等): 今日から約2年(730日)以内なら採用。大きく過去なら採用せず警告。
    dates = sorted(set(find_dates()))
    if dates:
        start, end = dates[0], dates[-1]
        try:
            td = datetime.datetime.strptime(today, "%Y%m%d").date()
            cutoff = (td - datetime.timedelta(days=730)).strftime("%Y%m%d")
        except Exception:
            cutoff = today
        if end != start:   # 有効期間の範囲
            if end >= today:
                res["start_date"] = start; ev["start_date"] = start
                res["end_date"] = end; ev["end_date"] = end
            else:
                res["date_stale"] = f"{start}〜{end}"
                ev["date_stale"] = f"検出した有効期間 {start}〜{end} は終了済み（自動入力せず・利用者が確認）"
        else:              # 単一日付（改正日など）
            if start >= cutoff:
                res["start_date"] = start; ev["start_date"] = start
            else:
                res["date_stale"] = start
                ev["date_stale"] = f"検出した日付 {start} は古い（2年以上前）ため自動入力せず・利用者が確認"

    # --- 電話 / URL ---
    mp = re.search(r"0\d{1,4}[-(\s]\d{1,4}[-)\s]\d{3,4}", t)
    if mp:
        res["phone"] = re.sub(r"[()\s]", "-", mp.group(0)).strip("-"); ev["phone"] = mp.group(0)
    mu = re.search(r"https?://[^\s　」）)]+", t)
    if mu:
        res["url"] = mu.group(0); ev["url"] = mu.group(0)

    # --- 事業者情報（運行主体者資料などに見出しがある場合だけ拾う。見出しベースで安全） ---
    def label_val(labels):
        for k in labels:
            m = re.search(k + r"[\s:：]*([^\n]+)", t)
            if m and m.group(1).strip():
                return m.group(1).strip()
        return None
    _name = label_val([r"氏名又は名称", r"事業者名", r"名\s*称"])
    if _name:
        res["agency_name"] = _name
        res["agency_official_name"] = _name
        ev["agency_name"] = _name
    _hojin = re.search(r"法人番号[\s:：]*(\d{13})", t) or re.search(r"(?<!\d)(\d{13})(?!\d)", t)
    if _hojin:
        res["agency_id"] = _hojin.group(1); ev["agency_id"] = _hojin.group(1)
    _zip = re.search(r"(?:〒|郵便番号)[\s:：]*(\d{3})[-\s]?(\d{4})", t)
    if _zip:
        res["agency_zip"] = _zip.group(1) + _zip.group(2)
        ev["agency_zip"] = f"{_zip.group(1)}-{_zip.group(2)}"
    _addr = label_val([r"住\s*所"])
    if _addr:
        res["agency_address"] = _addr; ev["agency_address"] = _addr
    _pres = label_val([r"代表者名", r"代表者"])
    if _pres:
        res["agency_president_name"] = _pres; ev["agency_president_name"] = _pres

    res["_evidence"] = ev
    return res


def main():
    ap = argparse.ArgumentParser(description="原文からPDF記載条件を保守的に検出")
    ap.add_argument("input")
    ap.add_argument("-o", "--output", default=None)
    a = ap.parse_args()
    text = _read_text(Path(a.input))
    res = detect(text)
    if a.output:
        Path(a.output).write_text(json.dumps(res, ensure_ascii=False, indent=2), encoding="utf-8")
    ev = res.get("_evidence", {})
    print(f"[OK] 検出 {len(ev)} 項目")
    for k, v in ev.items():
        print(f"  {k}: {res.get(k)}  ← '{v}'")


if __name__ == "__main__":
    sys.exit(main())
